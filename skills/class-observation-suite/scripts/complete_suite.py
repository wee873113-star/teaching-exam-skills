import os
import re
import sys
import argparse
import docx
from docx.shared import Inches, Pt
from docx.oxml.ns import qn

def offset_roc_date(roc_date_str, days_offset):
    digits = re.findall(r'\d+', roc_date_str)
    if len(digits) >= 3:
        roc_yr = int(digits[0])
        month = int(digits[1])
        day = int(digits[2])
        greg_yr = roc_yr + 1911
        from datetime import datetime, timedelta
        dt = datetime(greg_yr, month, day)
        new_dt = dt + timedelta(days=days_offset)
        new_roc_yr = new_dt.year - 1911
        return f"{new_roc_yr}年{new_dt.month}月{new_dt.day}日"
    return roc_date_str

def apply_font_to_paragraph(p, font_name, font_size):
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    for run in p.runs:
        run.font.name = font_name
        run.font.size = Pt(font_size)
        rPr = run._r.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn('w:eastAsia'), font_name)

def update_cell_text(cell, text, font_name, font_size):
    p0 = cell.paragraphs[0]
    p0.text = ""
    lines = text.split('\n')
    p0.text = lines[0]
    apply_font_to_paragraph(p0, font_name, font_size)
    
    for line in lines[1:]:
        p = cell.add_paragraph(line)
        apply_font_to_paragraph(p, font_name, font_size)

def find_file_by_pattern(directory, pattern):
    for f in os.listdir(directory):
        if re.search(pattern, f) and f.endswith('.docx') and not f.startswith('~$'):
            return os.path.join(directory, f)
    return None

def process_photos_file(template_path, photos_dir, photo_width, output_path, font_name="微軟正黑體", font_size=12, search_width=False):
    all_files = os.listdir(photos_dir)
    photo_files = [f for f in all_files if f.lower().endswith(('.jpg', '.jpeg', '.png'))]
    photo_files.sort()
    
    best_width = photo_width
    temp_doc_path = os.path.join(os.environ.get("TEMP", "C:\\Windows\\Temp"), "search_temp_suite.docx")
    
    if search_width:
        try:
            import win32com.client
            word_app = win32com.client.Dispatch("Word.Application")
            word_app.Visible = False
            
            print("Starting COM-based optimal photo width search...")
            widths_to_test = [round(2.5 + 0.05 * j, 2) for j in range(17)] # 2.50, 2.55, ..., 3.30
            for w in widths_to_test:
                write_photo_doc(template_path, photo_files, photos_dir, w, temp_doc_path, font_name, font_size)
                
                # Check page count
                doc = word_app.Documents.Open(
                    FileName=temp_doc_path,
                    ConfirmConversions=False,
                    ReadOnly=True,
                    AddToRecentFiles=False,
                    Visible=False
                )
                doc.Repaginate()
                pages = doc.ComputeStatistics(2)
                doc.Close(SaveChanges=False)
                
                print(f"Testing photo width: {w} inches -> Page Count: {pages}")
                if pages == 1:
                    best_width = w
                else:
                    print(f"Width {w} inches exceeds 1 page. Selecting best width {best_width} inches.")
                    break
            word_app.Quit()
        except Exception as e:
            print(f"Word COM application not available: {e}. Using default width: {photo_width} inches.")
        finally:
            if os.path.exists(temp_doc_path):
                try:
                    os.remove(temp_doc_path)
                except:
                    pass

    # Save final photo document
    write_photo_doc(template_path, photo_files, photos_dir, best_width, output_path, font_name, font_size)
    print(f"Saved: {output_path} (width: {best_width} inches)")

def write_photo_doc(template_path, photo_files, photos_dir, w, save_path, font_name, font_size):
    doc = docx.Document(template_path)
    table = doc.tables[0]
    
    for i in range(min(len(photo_files), 6)):
        photo_name = photo_files[i]
        photo_path = os.path.join(photos_dir, photo_name)
        desc_text = os.path.splitext(photo_name)[0]
        
        row_p = (i // 2) * 2
        col = i % 2
        row_d = row_p + 1

        # Photo cell
        photo_cell = table.rows[row_p].cells[col]
        p_photo = photo_cell.paragraphs[0]
        p_photo.text = ""
        p_photo.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
        p_photo.paragraph_format.space_before = Pt(0)
        p_photo.paragraph_format.space_after = Pt(0)
        p_photo.paragraph_format.line_spacing = 1.0

        run_photo = p_photo.add_run()
        run_photo.add_picture(photo_path, width=Inches(w))

        # Description cell
        desc_cell = table.rows[row_d].cells[col]
        p_desc = desc_cell.paragraphs[0]
        p_desc.text = ""
        p_desc.paragraph_format.space_before = Pt(0)
        p_desc.paragraph_format.space_after = Pt(0)
        p_desc.paragraph_format.line_spacing = 1.0
        
        run_desc_prefix = p_desc.add_run("說明：")
        run_desc_prefix.bold = True
        run_desc_prefix.font.name = font_name
        rPr = run_desc_prefix._r.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn('w:eastAsia'), font_name)
        
        run_desc_content = p_desc.add_run(desc_text)
        run_desc_content.font.name = font_name
        rPr2 = run_desc_content._r.get_or_add_rPr()
        rFonts2 = rPr2.get_or_add_rFonts()
        rFonts2.set(qn('w:eastAsia'), font_name)

    # Spacing optimization for trailing paragraph
    if len(doc.paragraphs) > 0:
        p_after = doc.paragraphs[-1]
        p_after.text = ""
        p_after.paragraph_format.space_before = Pt(0)
        p_after.paragraph_format.space_after = Pt(0)
        p_after.paragraph_format.line_spacing = Pt(1)
        run_after = p_after.add_run()
        run_after.font.size = Pt(1)
        
    doc.save(save_path)

def main():
    parser = argparse.ArgumentParser(description="觀議課全套表單自動化與格式化整合工具")
    parser.add_argument("--dir", required=True, help="包含觀議課檔案與照片的資料夾路徑")
    parser.add_argument("--font", default="微軟正黑體", help="填入欄位字型")
    parser.add_argument("--size", type=int, default=12, help="填入欄位字型大小")
    parser.add_argument("--photo-width", type=float, default=3.3, help="預設照片寬度（吋，設為3.3為極限一頁式）")
    parser.add_argument("--search-width", action="store_true", help="啟用 Word COM 引擎動態搜尋照片一頁式最佳寬度")
    
    group = parser.add_mutually_exclusive_group()
    group.add_argument("--only-forms", action="store_true", help="只填寫文檔表單 (不包含照片)")
    group.add_argument("--only-photos", action="store_true", help="只填寫照片文檔 (不包含表單)")
    
    args = parser.parse_args()
    p_dir = args.dir
    font_name = args.font
    font_size = args.size
    photo_width = args.photo_width
    search_width = args.search_width
    
    output_dir = os.path.join(p_dir, "完成檔")
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
        print(f"Created output directory: {output_dir}")

    # Find files dynamically
    activities_path = find_file_by_pattern(p_dir, r"活動設計單")
    notes_path = find_file_by_pattern(p_dir, r"觀課(紀錄|記錄)")
    pre_doc_path = find_file_by_pattern(p_dir, r"觀察前會談(紀錄|記錄)表")
    post_doc_path = find_file_by_pattern(p_dir, r"觀察後回饋會談(紀錄|記錄)表")
    eval_doc_path = find_file_by_pattern(p_dir, r"授課教師自評表")
    photos_doc_path = find_file_by_pattern(p_dir, r"教師同儕學習活動照片")
    photos_dir = os.path.join(p_dir, "照片")
    
    obs_doc_path = None
    for f in os.listdir(p_dir):
        if re.search(r"觀察(紀錄|記錄)表", f) and not re.search(r"(前|後)", f) and f.endswith('.docx') and not f.startswith('~$'):
            obs_doc_path = os.path.join(p_dir, f)
            break

    # Parse metadata if forms are needed
    if not args.only-photos:
        if not all([activities_path, notes_path, pre_doc_path, obs_doc_path, post_doc_path, eval_doc_path]):
            print("Error: Missing required documents in the directory for forms completion!")
            sys.exit(1)
            
        doc_act = docx.Document(activities_path)
        teaching_activities_text = doc_act.tables[1].rows[0].cells[0].text.strip()
        
        doc_notes = docx.Document(notes_path)
        meta_text = doc_notes.tables[0].rows[0].cells[0].text
        
        teacher = "吳宏文"
        grade = "三年級"
        subject = "健康與體育"
        observers = "陳皇成、王寓羚"
        unit = "籃球運轉手"
        obs_date_str = "114年11月14日"
        
        m_teacher = re.search(r'授課教師：\s*_*([^\s_]+)', meta_text)
        if m_teacher:
            teacher = m_teacher.group(1).strip()
        m_grade = re.search(r'任教年級：\s*_*([^\s_]+)', meta_text)
        if m_grade:
            grade = m_grade.group(1).strip()
        m_observers = re.search(r'回饋人員：\s*_*([^\s_]+)', meta_text)
        if m_observers:
            observers = m_observers.group(1).strip('_').strip()
        m_unit = re.search(r'教學單元：\s*_*([^\s_]+)', meta_text)
        if m_unit:
            unit = m_unit.group(1).strip()
        m_date = re.search(r'觀察日期：\s*_*(\d+)_*年\s*_*(\d+)_*月\s*_*(\d+)_*日', meta_text)
        if m_date:
            obs_date_str = f"{m_date.group(1)}年{m_date.group(2)}月{m_date.group(3)}日"
            
        pre_date_str = offset_roc_date(obs_date_str, -7)
        post_date_str = offset_roc_date(obs_date_str, 7)
        obs_location = "室外綜合球場"
        
        print("Successfully loaded metadata and dates.")

        # ==========================================
        # Fill 1: 觀察前會談紀錄表
        # ==========================================
        print("Completing 觀察前會談紀錄表...")
        doc_pre = docx.Document(pre_doc_path)
        t_pre = doc_pre.tables[0]
        update_cell_text(t_pre.rows[0].cells[1], teacher, font_name, font_size)
        update_cell_text(t_pre.rows[0].cells[4], grade, font_name, font_size)
        update_cell_text(t_pre.rows[0].cells[6], subject, font_name, font_size)
        update_cell_text(t_pre.rows[1].cells[1], observers, font_name, font_size)
        update_cell_text(t_pre.rows[2].cells[5], unit, font_name, font_size)
        update_cell_text(t_pre.rows[3].cells[1], pre_date_str, font_name, font_size)
        update_cell_text(t_pre.rows[4].cells[0], teaching_activities_text, font_name, font_size)
        
        p_r5 = t_pre.rows[5].cells[0].paragraphs
        p_r5[8].text = f"公開授課日期： {obs_date_str}"
        apply_font_to_paragraph(p_r5[8], font_name, font_size)
        p_r5[9].text = f"地點： {obs_location}"
        apply_font_to_paragraph(p_r5[9], font_name, font_size)
        p_r5[11].text = f"回饋會談預定日期： {post_date_str}"
        apply_font_to_paragraph(p_r5[11], font_name, font_size)
        
        pre_out = os.path.join(output_dir, os.path.basename(pre_doc_path))
        doc_pre.save(pre_out)
        print(f"Saved: {pre_out}")

        # ==========================================
        # Fill 2: 觀察紀錄表
        # ==========================================
        print("Completing 觀察紀錄表...")
        doc_obs = docx.Document(obs_doc_path)
        t_obs = doc_obs.tables[0]
        update_cell_text(t_obs.rows[0].cells[1], teacher, font_name, font_size)
        update_cell_text(t_obs.rows[0].cells[3], grade, font_name, font_size)
        update_cell_text(t_obs.rows[0].cells[5], subject, font_name, font_size)
        update_cell_text(t_obs.rows[1].cells[1], observers, font_name, font_size)
        update_cell_text(t_obs.rows[2].cells[1], unit, font_name, font_size)
        update_cell_text(t_obs.rows[3].cells[1], obs_date_str, font_name, font_size)
        update_cell_text(t_obs.rows[3].cells[4], obs_location, font_name, font_size)
        
        obs_out = os.path.join(output_dir, os.path.basename(obs_doc_path))
        doc_obs.save(obs_out)
        print(f"Saved: {obs_out}")

        # ==========================================
        # Fill 3: 觀察後回饋會談紀錄表
        # ==========================================
        print("Completing 觀察後回饋會談紀錄表...")
        doc_post = docx.Document(post_doc_path)
        t_post = doc_post.tables[0]
        update_cell_text(t_post.rows[0].cells[1], teacher, font_name, font_size)
        update_cell_text(t_post.rows[0].cells[3], grade, font_name, font_size)
        update_cell_text(t_post.rows[0].cells[5], subject, font_name, font_size)
        update_cell_text(t_post.rows[1].cells[1], observers, font_name, font_size)
        update_cell_text(t_post.rows[2].cells[1], unit, font_name, font_size)
        update_cell_text(t_post.rows[3].cells[1], post_date_str, font_name, font_size)
        
        strengths = [
            "教師設計的「籃球運轉手」單元，能有效融合運動技巧與遊戲互動，連結學生新舊知能，引發並維持學習動機。",
            "教師能清楚說明活動規則並示範運球技巧，亦會示範錯誤動作協助學生注意，提供適當的練習活動，學生參與度高。",
            "溝通技巧良好，音量與速度適中，能以口語提問、肢體語言及走動引導思考與行動，並以「很好！」等正向語言給予學生鼓勵。",
            "運用多元評量方式（觀察、口頭提問、操作回饋等），能針對學生學習表現即時給予調整指導與姿勢調整，幫助學生動作之精進。"
        ]
        strengths_text = "優點方面\n" + "\n".join([f"• {s}" for s in strengths])
        update_cell_text(t_post.rows[4].cells[0], strengths_text, font_name, font_size)
        
        original_share_title = t_post.rows[5].cells[0].paragraphs[0].text.strip()
        share_content = (
            "「本次公開授課讓教學者與觀課者皆獲益良多。透過同儕專業對話，教學者能更精準觀察學生在籃球運動中的團隊合作與互動表現；"
            "觀課者亦從中學習到如何有效結合遊戲與運動技巧，以及多元評量的具體實作，對未來體育教學中引導學生自主學習與同儕合作深具啟發。」"
        )
        update_cell_text(t_post.rows[5].cells[0], original_share_title + "\n" + share_content, font_name, font_size)
        
        original_action_title = t_post.rows[5].cells[2].paragraphs[0].text.strip()
        action_content = (
            "「1. 在未來的體育課程中，持續加強學生防守技能與護球動作的整合引導。 "
            "2. 持續使用口語與非口語指令（如吹哨與分組機制），提升課堂常規流暢度。 "
            "3. 下次觀察焦點預計安排在『學生同儕間的團隊策略討論與分工執行』，並採用同儕互評工具進行觀察紀錄。」"
        )
        update_cell_text(t_post.rows[5].cells[2], original_action_title + "\n" + action_content, font_name, font_size)
        
        post_out = os.path.join(output_dir, os.path.basename(post_doc_path))
        doc_post.save(post_out)
        print(f"Saved: {post_out}")

        # ==========================================
        # Fill 4: 授課教師自評表
        # ==========================================
        print("Completing 授課教師自評表...")
        doc_eval = docx.Document(eval_doc_path)
        t_eval = doc_eval.tables[0]
        
        update_cell_text(t_eval.rows[0].cells[1], observers, font_name, font_size)
        update_cell_text(t_eval.rows[0].cells[4], obs_date_str, font_name, font_size)
        update_cell_text(t_eval.rows[1].cells[1], teacher, font_name, font_size)
        update_cell_text(t_eval.rows[1].cells[4], grade, font_name, font_size)
        update_cell_text(t_eval.rows[2].cells[1], f"{subject} / {unit}", font_name, font_size)
        
        teaching_activities = (
            "本節公開授課單元為『籃球運轉手』。教學活動分為三階段：\n"
            "1. 準備活動：確認場地安全性，進行身體熱身，使用樂樂棒球進行繞身球感練習，以及原地左右手低中高運球與跑動運球熱身。\n"
            "2. 發展活動：實施『運球大作戰』，透過示範運球技巧與錯誤動作，引導學生在方框中運球並拍走他人球；隨後實施『運球攻城堡』，以分組（每組4人）方式運球繞過彩色角錐完成攻城遊戲，並融入合作與競爭策略。\n"
            "3. 綜合活動：引導學生進行分享討論，並由教師進行動作總結。"
        )
        student_performance = (
            "1. 學生在準備活動中積極參與，熱身運球時專注度高。\n"
            "2. 在『運球大作戰』與『運球攻城堡』等遊戲中，學生展現了極高的活動參與度、學習動機與團隊合作精神，能遵守遊戲規範，並與同儕進行友善的互動與競爭。\n"
            "3. 在綜合活動中，學生能認真回顧，勇於發表分享與討論，並在評量中展現出良好的運球、護球與傳接球動作技能。"
        )
        objective_achievement = (
            "本節課的學習目標達成情形如下：\n"
            "1. 學生已能認識籃球基本動作要領（如運球與護球的姿勢），並在練習中實踐。\n"
            "2. 透過分組遊戲，學生展現了優良的團隊合作精神，並與同學維持友善的互動關係，達成率達95%以上。\n"
            "3. 學生在實際運球跑動與攻城活動中，能具體表現出運球、護球與傳接球的動作技能。\n"
            "4. 學生能順利運用合作與競爭策略完成『運球大作戰』與『運球攻城堡』活動，整體學習目標圓滿達成。"
        )
        self_reflection = (
            "1. 本次教學設計結合了遊戲化機制，成功提升了學生的學習動機與參與度。學生在輕鬆有趣的環境下，能主動練習籃球的運球與護球技巧，達到寓教於樂的效果。\n"
            "2. 在教學過程中，除了示範正確運球技巧外，特別加入錯誤動作示範，有效幫助學生預防常見的姿勢錯誤，提升了學習成效。\n"
            "3. 多元評量的即時引導與正向語言鼓勵（如『很好！』），對於建立學生的自信心有顯著成效。\n"
            "4. 未來教學中，可進一步針對不同能力的學生設計分層難度的運球路線，並給予更細緻的個別化指導，以落實因材施教。"
        )
        feedback_reflections = (
            "感謝觀課同儕陳皇成與王寓羚老師的寶貴回饋。兩位老師肯定了本次『籃球運轉手』教學中將運動技巧融入遊戲的設計，並讚許學生的高度專注力與合作行為。在同儕回饋中，教學者也體會到口語與非口語指令（如吹哨與分組機制）對維持課堂常規流暢度的重要作用。\n"
            "針對同儕的建議，未來將持續加強護球與防守動作的實務指導，並計畫於下次公開授課時，將觀察焦點放在『學生同儕間的團隊策略討論與分工執行』，以期藉由同儕專業對話，不斷精進自我教學，實踐教與學的雙向成長。"
        )
        
        update_cell_text(t_eval.rows[4].cells[1], teaching_activities, font_name, font_size)
        update_cell_text(t_eval.rows[4].cells[3], student_performance, font_name, font_size)
        update_cell_text(t_eval.rows[5].cells[1], objective_achievement, font_name, font_size)
        update_cell_text(t_eval.rows[6].cells[1], self_reflection, font_name, font_size)
        update_cell_text(t_eval.rows[7].cells[1], feedback_reflections, font_name, font_size)
        
        eval_out = os.path.join(output_dir, os.path.basename(eval_doc_path))
        doc_eval.save(eval_out)
        print(f"Saved: {eval_out}")

    # ==========================================
    # Fill 5: 教師同儕學習活動照片
    # ==========================================
    if not args.only-forms:
        if not all([photos_doc_path, os.path.exists(photos_dir)]):
            print("Error: Missing required photos document or photos folder!")
            sys.exit(1)
            
        print("Completing 教師同儕學習活動照片 (One-Page layout optimization)...")
        photo_out = os.path.join(output_dir, os.path.basename(photos_doc_path))
        process_photos_file(photos_doc_path, photos_dir, photo_width, photo_out, font_name, font_size, search_width)

    print("\nAll requested tasks completed successfully!")

if __name__ == "__main__":
    main()
