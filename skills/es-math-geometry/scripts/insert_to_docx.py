#!/usr/bin/env python3
"""
insert_to_docx.py — 將幾何圖形（PNG）插入 Word 文件

Usage (library - 供其他技能呼叫):
  from insert_to_docx import insert_figure, insert_figures_side_by_side

Usage (CLI):
  python3 insert_to_docx.py <manifest.json> <target.docx> [--output out.docx] [--width 6]
"""
from pathlib import Path
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json, sys


def insert_figure(doc, png_path, width_cm=7.0, caption=None, align='center'):
    """
    在 Word 文件中插入單張幾何圖形。
    回傳插入的段落物件。
    """
    para = doc.add_paragraph()
    para.alignment = (WD_ALIGN_PARAGRAPH.CENTER
                      if align == 'center'
                      else WD_ALIGN_PARAGRAPH.LEFT)
    run = para.add_run()
    run.add_picture(str(png_path), width=Cm(width_cm))

    if caption:
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in cap.runs:
            r.font.size = Pt(9)
            r.font.italic = True
    return para


def insert_figures_side_by_side(doc, png_paths, width_cm=5.0, captions=None):
    """
    將多張圖形並排插入（使用無框表格）。
    png_paths: list of str/Path
    captions: list of str（可省略）
    """
    n = len(png_paths)
    table = doc.add_table(rows=2 if captions else 1, cols=n)
    table.style = 'Table Grid'
    # 移除框線（視覺上並排）
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tblPr = table._tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for border in ['top','left','bottom','right','insideH','insideV']:
        el = OxmlElement(f'w:{border}')
        el.set(qn('w:val'), 'none')
        tblBorders.append(el)
    tblPr.append(tblBorders)

    for i, png in enumerate(png_paths):
        cell = table.rows[0].cells[i]
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run()
        run.add_picture(str(png), width=Cm(width_cm))

    if captions:
        for i, cap in enumerate(captions):
            if i < n and cap:
                cell = table.rows[1].cells[i]
                para = cell.paragraphs[0]
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = para.add_run(cap)
                run.font.size = Pt(9)
                run.font.italic = True
    return table


def figures_from_manifest(manifest_path, doc, width_cm=7.0, mode='single'):
    """
    依 manifest.json 批次插入圖形。
    mode: 'single'（逐一插入）或 'row'（所有圖並排）
    """
    with open(manifest_path, 'r', encoding='utf-8') as f:
        manifest = json.load(f)

    png_paths = [item['png'] for item in manifest if 'png' in item]
    captions  = [item.get('caption','') for item in manifest if 'png' in item]

    if mode == 'row':
        insert_figures_side_by_side(doc, png_paths, width_cm=width_cm, captions=captions)
    else:
        for png, cap in zip(png_paths, captions):
            insert_figure(doc, png, width_cm=width_cm, caption=cap or None)


if __name__ == '__main__':
    import argparse
    parser = argparse.ArgumentParser(description='插入幾何圖形到 Word')
    parser.add_argument('manifest', help='manifest.json 路徑')
    parser.add_argument('docx',     help='目標 .docx 路徑')
    parser.add_argument('--output', help='輸出路徑（預設覆蓋原檔）')
    parser.add_argument('--width',  type=float, default=7.0, help='圖片寬度(cm)')
    parser.add_argument('--mode',   default='single', choices=['single','row'])
    args = parser.parse_args()

    target = Path(args.docx)
    doc = Document(str(target)) if target.exists() else Document()
    figures_from_manifest(args.manifest, doc, width_cm=args.width, mode=args.mode)
    out = args.output or str(target)
    doc.save(out)
    print(f"✅ 已儲存：{out}")
